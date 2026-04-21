"""
/**
 * @vibe-intent 底层文档渲染引擎：将标准MD结构及Graphviz配置转换为软著验收标准的Word文档 (AST 模板重构版)
 * @vibe-model Gemini 3.1 Pro
 * @vibe-ref intents.md#2026-04-21
 */
"""
import os
import re
import sys
import mistletoe
from mistletoe.block_token import Heading, Paragraph, List, ListItem, BlockCode
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw

try:
    import graphviz
    from graphviz import Digraph
    HAS_GRAPHVIZ = True
except ImportError:
    HAS_GRAPHVIZ = False

# --- 图表生成相关逻辑 ---
def parse_flow_config(md_text):
    pattern_dot = re.compile(r'\[FLOW-CONFIG-START\](.*?)\[FLOW-CONFIG-END\]', re.DOTALL)
    match = pattern_dot.search(md_text)
    if not match: 
        return None, None
        
    content = match.group(1).strip()
    if 'digraph' in content:
        return 'dot', content
    
    nodes, edges = {}, []
    lines = content.split('\n')
    current_section = None
    
    for line in lines:
        line = line.strip()
        if not line: continue
        if '[Nodes]' in line: current_section = 'nodes'; continue
        elif '[Edges]' in line: current_section = 'edges'; continue
        
        if current_section == 'nodes':
            parts = line.split(':', 1)
            if len(parts) >= 2:
                n_id = parts[0].strip()
                attrs = parts[1].split('|')
                label = attrs[0].strip()
                shape = 'box'
                if len(attrs) > 1 and 'shape=' in attrs[1]: shape = attrs[1].split('=')[1].strip()
                nodes[n_id] = {'label': label, 'shape': shape}
        elif current_section == 'edges' and '->' in line:
            s, d = line.split('->')
            edges.append((s.strip(), d.strip()))
    return nodes, edges

def draw_dynamic_flowchart(nodes, edges, filename='temp_flowchart.png', dot_code=None):
    if not HAS_GRAPHVIZ: return False
    try:
        from graphviz import Source
        if dot_code:
            src = Source(dot_code)
            src.render(outfile=filename, format='png', cleanup=True)
            return True
            
        dot = Digraph(comment='Auto Flow', format='png')
        dot.attr(rankdir='TB', size='6,8')
        dot.attr('node', fontname='SimHei', fontsize='12', style='filled', fillcolor='aliceblue')
        dot.attr('edge', fontname='SimHei')
        for n_id, props in nodes.items():
            fill = 'lightgrey' if props['shape'] == 'oval' else 'aliceblue'
            if props['shape'] == 'diamond': fill = 'lightyellow'
            dot.node(n_id, props['label'], shape=props['shape'], fillcolor=fill)
        for src, dst in edges: dot.edge(src, dst)
        dot.render(filename.replace('.png', ''), cleanup=True)
        return True
    except Exception as e: 
        print(f"Graphviz draw failed: {e}")
        return False

def create_dummy_image(path):
    img = Image.new('RGB', (600, 350), (240, 240, 240))
    d = ImageDraw.Draw(img); d.rectangle([0,0,599,349], outline="gray")
    img.save(path)

class ASTRenderer:
    def __init__(self, template_path=None, software_name="软件说明书", max_heading_level=3, flow_ready=False):
        if template_path and os.path.exists(template_path):
            self.doc = Document(template_path)
            # 对模板自带的段落不用清理，我们只需追加即可。如果有特殊需求，此处可以改进为先清空
            # 如果《范例》文档里面自带第一页等内容，保留它是符合“范例结构”的，我们生成的内容追加。
            # 或者按照要求直接保存。为了避免与《范例说明》的正文冲突，可以在 docx 中预留书签插入，
            # 但目前要求直接追加内容，我们会追加到后方。
        else:
            print("警告：未找到模板文件，将使用默认样式。")
            self.doc = Document()
            
        self.software_name = software_name
        self.max_level = max_heading_level
        self.flow_ready = flow_ready
        
        self.dummy_img_path = "temp_placeholder.png"
        self.flow_img_path = "temp_flowchart.png"
        if not os.path.exists(self.dummy_img_path): 
            create_dummy_image(self.dummy_img_path)

    def render(self, markdown_text, output_path):
        doc_ast = mistletoe.Document(markdown_text)
        for node in doc_ast.children:
            self._visit(node)
        self.doc.save(output_path)
        
        for f in [self.dummy_img_path, self.flow_img_path]: 
            if os.path.exists(f): os.remove(f)

    def _visit(self, node):
        if isinstance(node, Heading):
            self._handle_heading(node)
        elif isinstance(node, Paragraph):
            self._handle_paragraph(node)
        elif isinstance(node, List):
            self._handle_list(node)
        elif isinstance(node, BlockCode):
            pass # 抛弃纯代码块

    def _handle_heading(self, node):
        text = self._extract_text(node)
        level = node.level
        if level > self.max_level:
            p = self.doc.add_paragraph(style='Normal')
            run = p.add_run(text)
            run.bold = True
        else:
            self.doc.add_heading(text, level=level)

    def _handle_paragraph(self, node):
        text = self._extract_text(node)
        if re.match(r'^\[[图圖]\d+[：:].*\]$', text):
            p_img = self.doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            target_img = self.flow_img_path if ("图1" in text and self.flow_ready) else self.dummy_img_path
            p_img.add_run().add_picture(target_img, width=Cm(14))
            
            p_cap = self.doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_cap.add_run(text.strip('[]'))
            run.font.size = Pt(10.5)
        elif text.strip() == "":
            pass
        else:
            p = self.doc.add_paragraph(text, style='Normal')
            p.paragraph_format.first_line_indent = Cm(0.85)

    def _handle_list(self, node):
        for item in node.children:
            if isinstance(item, ListItem):
                text = self._extract_text(item)
                try:
                    self.doc.add_paragraph(text, style='List Bullet')
                except KeyError:
                    # 兼容部分中文模板没有 'List Bullet' 样式的情况
                    p = self.doc.add_paragraph(f"• {text}", style='Normal')
                    p.paragraph_format.left_indent = Cm(0.7)

    def _extract_text(self, node):
        text = ""
        if hasattr(node, 'children'):
            for child in node.children:
                if hasattr(child, 'content'):
                    text += child.content
                else:
                    text += self._extract_text(child)
        elif hasattr(node, 'content'):
            text += node.content
        return text

# --- 主程序 ---
def generate_docx(md_path, docx_path, software_name, template_path):
    with open(md_path, 'r', encoding='utf-8') as f: 
        content = f.read()

    nodes, edges = parse_flow_config(content)
    flow_ready = False
    
    if nodes == 'dot':
        flow_ready = draw_dynamic_flowchart(None, None, dot_code=edges)
    elif nodes and edges:
        flow_ready = draw_dynamic_flowchart(nodes, edges)
    
    content = re.sub(r'\[FLOW-CONFIG-START\][\s\S]*?\[FLOW-CONFIG-END\]', '', content)

    renderer = ASTRenderer(template_path=template_path, software_name=software_name, flow_ready=flow_ready, max_heading_level=3)
    renderer.render(content, docx_path)
    print(f"[Success] Word 文档已生成: {docx_path}")

if __name__ == "__main__":
    if len(sys.argv) < 5:
        print("Usage: python universal_doc_gen.py <md_path> <docx_path> <software_name> <template_path>")
        sys.exit(1)
    generate_docx(sys.argv[1], sys.argv[2], sys.argv[3], sys.argv[4])
